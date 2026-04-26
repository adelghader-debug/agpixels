"""Generate AGPixels Google Ads API tool design document (DOCX)."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x0E, 0x1A)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = k
        c1.text = v
        for para in c0.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(11)
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
    return table


# -----------------------------------------------------------------------------

doc = Document()

# Set default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AGPixels Google Ads API Tool")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x6B, 0x5D, 0xF7)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Design Documentation")
run.font.size = Pt(14)
run.italic = True
run.font.color.rgb = RGBColor(0x6B, 0x6A, 0x78)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Prepared for Google Ads API Basic Access Application\nApril 26, 2026")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x6B, 0x6A, 0x78)

doc.add_paragraph()

# 1. Overview
add_heading(doc, "1. Tool Overview", level=1)
add_para(
    doc,
    "AGPixels is a small Canadian web development and digital agency operating "
    "at https://agpixels.ca. We build and maintain websites and mobile applications "
    "for our own brand and for client businesses across healthcare, dog training, "
    "and senior care sectors.",
)
add_para(
    doc,
    "The AGPixels Google Ads Tool is an internal tool used exclusively by the "
    "AGPixels team to manage and report on Google Ads campaigns across two categories "
    "of accounts:",
)
add_bullet(doc, "AGPixels' own advertising account, used to promote our web design, "
                "maintenance, and mobile app development services.")
add_bullet(doc, "Client advertising accounts where AGPixels has been granted manager "
                "access via the standard Google Ads invitation flow.")

add_para(
    doc,
    "The tool is not offered as a third-party service, is not user-facing, and is not "
    "marketed externally. It is an internal operations and reporting layer for our "
    "agency.",
)

# 2. Architecture
add_heading(doc, "2. Architecture", level=1)
add_para(doc, "The tool consists of two layers:")
add_para(doc, "Backend (Cloudflare Workers + Python scripts)", bold=True)
add_bullet(doc, "Cloudflare Workers handle scheduled and on-demand API calls.")
add_bullet(doc, "Local Python scripts (using google-ads-python) are used by the "
                "AGPixels team for ad-hoc campaign management and reporting.")
add_bullet(doc, "All API calls authenticate using OAuth 2.0 refresh tokens stored "
                "securely as Cloudflare Worker secrets / environment variables — "
                "never committed to source control.")

add_para(doc, "Reporting layer", bold=True)
add_bullet(doc, "Performance reports are generated as PDF/CSV exports and shared "
                "with clients via email.")
add_bullet(doc, "Reports may be cached briefly (under 24 hours) for performance.")

add_kv_table(doc, [
    ("Hosting", "Cloudflare (Canada)"),
    ("Languages", "Python (server-side scripts), JavaScript (Worker)"),
    ("API library", "google-ads-python (official Google Ads client library)"),
    ("Storage", "Minimal — credentials in env vars, ephemeral query results only"),
    ("Authentication", "OAuth 2.0 with refresh tokens per managed account"),
])

# 3. Google Ads API usage
add_heading(doc, "3. Google Ads API Usage", level=1)
add_para(doc, "The tool calls the following Google Ads API services:")

add_para(doc, "Read operations (queries via GoogleAdsService.SearchStream)", bold=True)
add_bullet(doc, "campaign — list and inspect campaigns, budgets, status.")
add_bullet(doc, "ad_group — list ad groups under each campaign.")
add_bullet(doc, "ad_group_ad — review active ads and ad assets.")
add_bullet(doc, "keyword_view — review keyword performance.")
add_bullet(doc, "metrics — pull impression, click, cost, conversion data for reporting.")
add_bullet(doc, "campaign_budget — check budget status and pacing.")

add_para(doc, "Write operations (mutations)", bold=True)
add_bullet(doc, "CampaignService.MutateCampaigns — create, pause, enable, rename campaigns.")
add_bullet(doc, "CampaignBudgetService.MutateCampaignBudgets — adjust daily budgets.")
add_bullet(doc, "AdGroupService.MutateAdGroups — manage ad groups.")
add_bullet(doc, "AdGroupCriterionService.MutateAdGroupCriteria — add/remove keywords, "
                "adjust bids, manage negative keywords.")
add_bullet(doc, "AdGroupAdService.MutateAdGroupAds — create and manage responsive search ads.")

add_para(doc, "Anticipated request volume", bold=True)
add_bullet(doc, "Less than 10,000 API operations per day across all managed accounts.")
add_bullet(doc, "Read operations: roughly 200–500/day for routine reporting.")
add_bullet(doc, "Write operations: typically under 50/day during normal operations.")

# 4. Data flow
add_heading(doc, "4. Data Flow", level=1)
add_para(doc, "1. AGPixels team member initiates an action (CLI, dashboard, or scheduled job).")
add_para(doc, "2. The tool authenticates to the Google Ads API using the OAuth refresh "
              "token for the target account, which the account owner has explicitly granted.")
add_para(doc, "3. The tool issues a GAQL query or mutation request.")
add_para(doc, "4. Response data is processed locally and either stored ephemerally for "
              "the duration of the session or written to a CSV/PDF report.")
add_para(doc, "5. Reports are delivered to AGPixels staff or, when applicable, to the "
              "client account owner.")

add_para(doc, "No Google Ads data is exposed publicly, embedded in third-party services, "
              "or used to train any external models.")

# 5. User access & permissions
add_heading(doc, "5. User Access and Permissions", level=1)
add_bullet(doc, "Only authorized AGPixels team members can run the tool. Access is "
                "controlled via local credentials and Cloudflare Worker secret bindings.")
add_bullet(doc, "Each managed Google Ads account requires a separate OAuth refresh token, "
                "obtained via the standard OAuth consent flow with the account owner "
                "explicitly approving access.")
add_bullet(doc, "Client accounts are added through the standard Google Ads manager-account "
                "invitation flow. Clients can revoke access at any time from their own "
                "Google Ads admin panel.")
add_bullet(doc, "The tool maintains an audit log of write operations for our own records.")

# 6. Security
add_heading(doc, "6. Security", level=1)
add_bullet(doc, "Developer token, OAuth client secret, and refresh tokens are stored as "
                "Cloudflare Worker secrets and / or in OS-level environment variables. "
                "They are never committed to source control.")
add_bullet(doc, "All API traffic uses HTTPS / TLS 1.2+.")
add_bullet(doc, "Access to the source code repository is restricted to the AGPixels team.")
add_bullet(doc, "Credentials are rotated when team members leave.")
add_bullet(doc, "We do not log raw API responses; only aggregated metrics and operation results.")

# 7. Error handling & rate limiting
add_heading(doc, "7. Error Handling and Rate Limiting", level=1)
add_bullet(doc, "The tool uses the official google-ads-python client, which provides "
                "built-in retry logic with exponential backoff for transient errors.")
add_bullet(doc, "Quota errors (RESOURCE_EXHAUSTED) trigger a backoff and notification "
                "to the AGPixels team rather than aggressive retry.")
add_bullet(doc, "All write operations are wrapped in try/except with meaningful logging "
                "so failures are surfaced and acted on, not silently retried.")

# 8. Compliance
add_heading(doc, "8. Compliance with Google Ads API Policy", level=1)
add_bullet(doc, "We have read and agree to the Google Ads API Terms and Conditions, the "
                "Google Ads API Required Minimum Functionality (RMF) policy, and the "
                "Google Ads Policies.")
add_bullet(doc, "We do not sell, rent, or transfer Google Ads data to third parties.")
add_bullet(doc, "We do not use the API to access accounts we have not been explicitly "
                "granted permission to manage.")
add_bullet(doc, "We will display required disclaimers (e.g. that data is provided by "
                "Google Ads) on any client-facing reports.")
add_bullet(doc, "We comply with the Personal Information Protection and Electronic "
                "Documents Act (PIPEDA) for handling Canadian client data.")

# 9. Future plans
add_heading(doc, "9. Future Plans", level=1)
add_para(
    doc,
    "Initially the tool will support Search and Performance Max campaigns, which are the "
    "primary campaign types we run for AGPixels and our clients. We may expand to Display "
    "campaigns as our client base grows. We do not currently plan to use App Conversion "
    "Tracking or the Remarketing API.",
)

# 10. Contact
add_heading(doc, "10. Contact", level=1)
add_kv_table(doc, [
    ("Company", "AGPixels"),
    ("Website", "https://agpixels.ca"),
    ("Contact email", "adel.ghader@gmail.com"),
    ("Manager account (MCC) ID", "641-583-9842"),
])

# Save
out = "agpixels-google-ads-tool-design.docx"
doc.save(out)
print(f"Saved: {out}")
